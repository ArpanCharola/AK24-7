"""Resume → structured profile parsing.

Takes an uploaded resume (PDF or DOCX), extracts its plain text, and asks the
shared LLM helper to project it into a structured profile draft that the user
reviews and edits before saving. Parsing is best-effort: nothing here raises —
on any failure we return a usable draft (empty structure + whatever raw text we
recovered) so the upload endpoint always responds with something editable.

The structured shape produced here is the contract the profile editor and the
tailored-resume engine both build on:

    {
      "contact": {full_name, email, phone, location,
                  linkedin_url, github_url, website_url},
      "summary": str,
      "work_experience": [{company, title, location, start_date, end_date,
                           bullets: [str]}],
      "education":       [{institution, degree, field, start_date, end_date, grade}],
      "skills":          [str],
      "projects":        [{name, description, tech: [str], url}],
      "certifications":  [{name, issuer, year}],
    }
"""
import io
import logging
import re

from app.services.resume_tailor import _generate, _parse_json

logger = logging.getLogger(__name__)

# Scalar contact fields map straight onto existing User columns. `email` is the
# login identity — we surface it in the draft but never overwrite it on save.
CONTACT_FIELDS = (
    "full_name", "email", "phone", "location",
    "linkedin_url", "github_url", "website_url",
)
# Structured sections persisted as JSON-in-Text columns on `users`
# (landed by Foundation's migration — see WORKSTREAMS #6).
PROFILE_SECTIONS = ("work_experience", "education", "skills", "projects", "certifications")

_COMMON_SKILLS = (
    "python", "javascript", "typescript", "java", "react", "next.js", "node.js",
    "express", "django", "fastapi", "flask", "sql", "postgresql", "mysql",
    "mongodb", "redis", "aws", "azure", "gcp", "docker", "kubernetes", "git",
    "html", "css", "tailwind", "machine learning", "deep learning", "nlp",
    "pandas", "numpy", "tensorflow", "pytorch", "openai", "langchain",
)
_COMMON_ROLES = (
    "Software Engineer", "Software Developer", "Frontend Developer",
    "Backend Developer", "Full Stack Developer", "Fullstack Developer",
    "React Developer", "Python Developer", "Java Developer", "AI Engineer",
    "Machine Learning Engineer", "Data Analyst", "Data Engineer",
)
_PHONE_RE = re.compile(r"(?:\+91[\s-]?)?[6-9]\d{9}\b")


def empty_profile() -> dict:
    """A fully-formed, empty draft — the shape every consumer can rely on."""
    return {
        "contact": {f: None for f in CONTACT_FIELDS},
        "summary": "",
        "work_experience": [],
        "education": [],
        "skills": [],
        "desired_roles": [],
        "projects": [],
        "certifications": [],
    }


def _clean_text(text: str | None) -> str:
    return re.sub(r"[ \t]+", " ", text or "").strip()


def _extract_pdf_text_pymupdf(data: bytes) -> str:
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        pages = []
        for page in doc:
            text = page.get_text("text", sort=True) or ""
            if text.strip():
                pages.append(text)
    return _clean_text("\n\n".join(pages))


def _extract_pdf_text_pdfplumber(data: bytes) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or page.extract_text(layout=True) or ""
            if text.strip():
                pages.append(text)
    return _clean_text("\n\n".join(pages))


def _extract_pdf_text_pypdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            pass
    pages = []
    for page in reader.pages:
        text = ""
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except TypeError:
            text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return _clean_text("\n\n".join(pages))


def _extract_pdf_text_ocr(data: bytes) -> str:
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf
    import pytesseract
    from PIL import Image

    pages = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for index, page in enumerate(doc):
            if index >= 4:
                break
            pix = page.get_pixmap(matrix=pymupdf.Matrix(3, 3), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
            image = image.point(lambda pixel: 255 if pixel > 180 else 0)
            text = pytesseract.image_to_string(image, config="--oem 3 --psm 6") or ""
            if text.strip():
                pages.append(text)
    return _clean_text("\n\n".join(pages))


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from digital PDFs using multiple libraries.

    Many resumes are generated by Canva/Word/LaTeX/export tools and one parser
    can fail while another succeeds. OCR runs only after digital extraction
    fails, covering scanned/image-only PDFs when Tesseract is installed.
    """
    for extract in (
        _extract_pdf_text_pymupdf,
        _extract_pdf_text_pdfplumber,
        _extract_pdf_text_pypdf,
        _extract_pdf_text_ocr,
    ):
        try:
            text = extract(data)
        except Exception as exc:
            logger.info("PDF extraction via %s failed: %s", extract.__name__, repr(exc)[:180])
            continue
        if text:
            return text
    return ""


def _extract_docx_text(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_upload(filename: str | None, data: bytes) -> str:
    """Extract plain text from a PDF or DOCX upload.

    Routes by extension, then falls back to trying both parsers so a
    mislabelled file still has a chance. Returns "" if nothing readable.
    """
    name = (filename or "").lower()
    extractors = []
    if name.endswith(".pdf"):
        extractors = [_extract_pdf_text, _extract_docx_text]
    elif name.endswith(".docx"):
        extractors = [_extract_docx_text, _extract_pdf_text]
    else:
        extractors = [_extract_pdf_text, _extract_docx_text]

    for extract in extractors:
        try:
            text = extract(data)
        except Exception as exc:
            logger.debug("Resume text extraction via %s failed: %s", extract.__name__, exc)
            continue
        if text:
            return text
    return ""


_PARSE_SYSTEM = (
    "You are a resume parser. Read the resume text and return a SINGLE JSON object "
    "with EXACTLY these keys:\n"
    '  "contact": {"full_name","email","phone","location","linkedin_url","github_url","website_url"},\n'
    '  "summary": string (the professional summary/objective, "" if none),\n'
    '  "work_experience": [{"company","title","location","start_date","end_date","bullets":[string]}],\n'
    '  "education": [{"institution","degree","field","start_date","end_date","grade"}],\n'
    '  "skills": [string],\n'
    '  "desired_roles": [string],\n'
    '  "projects": [{"name","description","tech":[string],"url"}],\n'
    '  "certifications": [{"name","issuer","year"}]\n'
    "Rules: use null for unknown scalar fields and [] for missing lists. Dates as written "
    "in the resume (e.g. 'Jan 2022', 'Present'). Split each role's responsibilities into "
    "separate bullet strings. Extract ONLY what is present — never invent employers, dates, "
    "skills, or metrics. Output only the JSON — no explanation, no markdown fences."
)


async def parse_resume_text(text: str) -> dict:
    """Project resume text into the structured draft. Never raises."""
    draft = _fallback_parse(text)
    text = (text or "").strip()
    if not text:
        return draft
    try:
        raw = await _generate(_PARSE_SYSTEM, f"Resume text:\n\n{text}", max_tokens=3000)
        parsed = _parse_json(raw)
    except Exception as exc:
        logger.warning("Resume LLM parse failed: %s", exc)
        return draft

    if not isinstance(parsed, dict) or (len(parsed) == 1 and "raw" in parsed):
        return draft
    parsed = _coerce_profile(parsed)
    parsed["contact"] = {**draft["contact"], **{k: v for k, v in parsed["contact"].items() if v}}
    parsed["skills"] = parsed["skills"] or draft["skills"]
    parsed["desired_roles"] = parsed["desired_roles"] or draft["desired_roles"]
    return parsed


def _fallback_parse(text: str | None) -> dict:
    out = empty_profile()
    text = text or ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if lines:
        first = lines[0]
        if 2 <= len(first.split()) <= 5 and not any(ch.isdigit() for ch in first):
            out["contact"]["full_name"] = first[:120]
    phone = _PHONE_RE.search(text)
    if phone:
        out["contact"]["phone"] = phone.group(0)
    low = text.lower()
    out["skills"] = [skill for skill in _COMMON_SKILLS if re.search(r"\b" + re.escape(skill) + r"\b", low)]
    out["desired_roles"] = [role for role in _COMMON_ROLES if role.lower() in low]
    return out


def _coerce_profile(parsed: dict) -> dict:
    """Normalise model output into the exact draft shape, dropping junk."""
    out = empty_profile()

    contact = parsed.get("contact")
    if isinstance(contact, dict):
        for f in CONTACT_FIELDS:
            v = contact.get(f)
            if isinstance(v, str):
                out["contact"][f] = v.strip() or None
            elif v is not None:
                out["contact"][f] = v

    summary = parsed.get("summary")
    out["summary"] = summary.strip() if isinstance(summary, str) else ""

    for key in ("work_experience", "education", "projects", "certifications"):
        v = parsed.get(key)
        if isinstance(v, list):
            out[key] = [item for item in v if isinstance(item, dict)]

    skills = parsed.get("skills")
    if isinstance(skills, list):
        out["skills"] = [s.strip() for s in skills if isinstance(s, str) and s.strip()]

    desired_roles = parsed.get("desired_roles")
    if isinstance(desired_roles, list):
        out["desired_roles"] = [r.strip() for r in desired_roles if isinstance(r, str) and r.strip()]

    return out


async def parse_resume_file(filename: str | None, data: bytes) -> dict:
    """Full pipeline: bytes → text → structured draft. Never raises.

    Returns the draft plus the recovered raw text under ``resume_text`` so the
    caller can persist it and show the user what was parsed against.
    """
    text = extract_text_from_upload(filename, data)
    draft = await parse_resume_text(text)
    draft["resume_text"] = text
    return draft


# ── Reading the saved profile back off the User row ───────────────────────────

def load_structured_profile(user) -> dict:
    """Reconstruct the structured draft from a User instance.

    Contact comes from scalar columns; the section lists come from the JSON
    columns landed by Foundation's migration. Defensive against the migration
    not having run yet (getattr default None) and against either JSON or
    JSON-encoded-Text storage.
    """
    import json

    out = empty_profile()
    for f in CONTACT_FIELDS:
        out["contact"][f] = getattr(user, f, None)

    summary = getattr(user, "summary", None)
    out["summary"] = summary or ""

    for key in PROFILE_SECTIONS:
        raw = getattr(user, key, None)
        if raw is None:
            continue
        if isinstance(raw, (list, dict)):
            value = raw
        elif isinstance(raw, str):
            try:
                value = json.loads(raw)
            except (TypeError, ValueError):
                continue
        else:
            continue
        if key == "skills":
            if isinstance(value, list):
                out[key] = [s for s in value if isinstance(s, str)]
        elif isinstance(value, list):
            out[key] = [item for item in value if isinstance(item, dict)]

    return out


def has_structured_profile(user) -> bool:
    """True if the user has saved any structured profile content."""
    profile = load_structured_profile(user)
    return bool(
        profile["summary"]
        or any(profile[s] for s in PROFILE_SECTIONS)
        or any(profile["contact"].values())
    )
